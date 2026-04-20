import os
import sys

def resource_path(relative_path):
    try:
        base_path = sys._MEIPASS
    except AttributeError:
        base_path = os.path.dirname(os.path.abspath(sys.argv[0]))
    return os.path.join(base_path, relative_path)
import json
import ctypes
ctypes.windll.shell32.SetCurrentProcessExplicitAppUserModelID("Ia_List_Creator")
import xml.etree.ElementTree as ET
from PyQt6.QtGui import QIcon
from PyQt6.QtWidgets import QMessageBox
from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from PyQt6.QtWidgets import (
    QApplication, QWidget, QVBoxLayout, QHBoxLayout, QPushButton,
    QListWidget, QFileDialog, QInputDialog, QLineEdit, QMenu, QDialog
)
from PyQt6.QtCore import Qt
from reportlab.lib.pagesizes import A4
from reportlab.pdfgen import canvas
from reportlab.lib.colors import HexColor, black, gray
from reportlab.lib.units import cm

DEFAULT_OBS = "Nenhuma observação adicionada."

def show_license_popup():
    msg = QMessageBox()
    msg.setWindowTitle("Aviso")
    msg.setIcon(QMessageBox.Icon.Information)
    msg.setText(
        "Programa criado e distribuído por: IA_List_Creatos\n\n"
        "O uso deste software é livre, o autor permite download, "
        "compartilhamento e qualquer tipo de uso não comercial.\n\n"
        "Se você pagou por ele foi roubado!"
    )
    msg.setStandardButtons(QMessageBox.StandardButton.Ok)
    msg.exec()

def is_hidden(filepath):
    if os.name == 'nt':
        try:
            attrs = ctypes.windll.kernel32.GetFileAttributesW(filepath)
            if attrs != -1:
                return bool(attrs & (0x2 | 0x4))
        except Exception:
            pass
    return os.path.basename(filepath).startswith('.')

def create_styled_button(text, color_hex):
    btn = QPushButton(text)
    btn.setMinimumHeight(45)
    btn.setMaximumHeight(45)
    btn.setStyleSheet(f"""
        QPushButton {{
            background-color: {color_hex};
            color: white;
            border: none;
            border-radius: 12px;
            padding: 0px 20px;
            font-size: 12px;
            font-weight: bold;
        }}
        QPushButton:hover {{
            background-color: {adjust_brightness(color_hex, -25)};
        }}
        QPushButton:pressed {{
            background-color: {adjust_brightness(color_hex, -50)};
        }}
    """)
    return btn

def adjust_brightness(hex_color, amount):
    hex_color = hex_color.lstrip('#')
    r, g, b = int(hex_color[0:2], 16), int(hex_color[2:4], 16), int(hex_color[4:6], 16)
    r = max(0, min(255, r + amount))
    g = max(0, min(255, g + amount))
    b = max(0, min(255, b + amount))
    return f"#{r:02x}{g:02x}{b:02x}"

class CollectionDialog(QDialog):
    def __init__(self, parent, full_order):
        super().__init__(parent)
        self.setWindowTitle("Modo Coleção")
        self.resize(600, 500)
        self.full_order = full_order
        self.filtered = []
        self.anchor_index = None

        layout = QVBoxLayout()
        self.search = QLineEdit()
        self.search.setPlaceholderText("Buscar...")
        self.search.textChanged.connect(self.filter)

        self.list = QListWidget()
        self.list.setDragDropMode(QListWidget.DragDropMode.InternalMove)
        self.list.setContextMenuPolicy(Qt.ContextMenuPolicy.CustomContextMenu)
        self.list.customContextMenuRequested.connect(self.menu)

        self.btn_save = QPushButton("Salvar coleção")
        self.btn_save.setStyleSheet("""
            QPushButton { background-color: #4CAF50; color: white; border: none; border-radius: 8px; padding: 10px; font-weight: bold; }
            QPushButton:hover { background-color: #45a049; }
        """)
        self.btn_save.clicked.connect(self.accept)

        layout.addWidget(self.search)
        layout.addWidget(self.list)
        layout.addWidget(self.btn_save)
        self.setLayout(layout)

    def filter(self):
        text = self.search.text().lower()
        self.list.clear()
        terms = [t.strip() for t in text.split(",") if t.strip()]
        self.filtered = self.full_order.copy() if not terms else [f for f in self.full_order if any(term in f.lower() for term in terms)]
        for f in self.filtered:
            idx = self.full_order.index(f) + 1
            self.list.addItem(f"{str(idx).zfill(4)}  - {os.path.splitext(f)[0]}")

    def menu(self, pos):
        menu = QMenu()
        act = menu.addAction("Definir âncora")
        action = menu.exec(self.list.mapToGlobal(pos))
        if action == act:
            self.anchor_index = self.list.currentRow()

    def get_new_order(self):
        if self.anchor_index is None:
            return None
        items = []
        for i in range(self.list.count()):
            text = self.list.item(i).text()
            name = text.split("  - ", 1)[1]
            for f in self.filtered:
                if os.path.splitext(f)[0] == name:
                    items.append(f)
        anchor_item = items[self.anchor_index]
        items.remove(anchor_item)
        items.insert(0, anchor_item)
        return items


class App(QWidget):
    def __init__(self):
        super().__init__()
        self.setWindowTitle("Gerador de Lista")
        self.setWindowIcon(QIcon(resource_path("icone.ico")))
        self.resize(900, 700)
        
        self.setStyleSheet("""
            QWidget { background-color: #2b2b2b; color: white; }
            QLineEdit { background-color: #3c3c3c; border: 2px solid #555555; border-radius: 8px; padding: 10px; font-size: 14px; color: white; }
            QLineEdit:focus { border: 2px solid #0078d4; }
        """)

        main_layout = QVBoxLayout()
        main_layout.setSpacing(10)
        main_layout.setContentsMargins(15, 15, 15, 15)

        self.search_main = QLineEdit()
        self.search_main.setPlaceholderText("Buscar na lista...")
        self.search_main.setMinimumHeight(40)
        self.search_main.setMaximumHeight(40)
        self.search_main.textChanged.connect(self.filter_main)
        main_layout.addWidget(self.search_main)

        self.list_widget = QListWidget()
        self.list_widget.keyPressEvent = self.handle_delete_key
        self.list_widget.setVerticalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAlwaysOn)
        self.list_widget.setHorizontalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAsNeeded)
        self.list_widget.setStyleSheet("""
            QListWidget { background-color: #3c3c3c; border: 2px solid #555555; border-radius: 8px; padding: 5px; font-size: 13px; }
            QListWidget::item { padding: 8px; border-radius: 4px; margin: 2px 0; border-bottom: 1px solid #4a4a4a; }
            QListWidget::item:selected { background-color: #0078d4; }
            QListWidget::item:hover:!selected { background-color: #4a4a4a; }
            QScrollBar:vertical { background: #2b2b2b; width: 12px; border-radius: 6px; }
            QScrollBar::handle:vertical { background: #666666; border-radius: 6px; min-height: 30px; }
            QScrollBar::handle:vertical:hover { background: #888888; }
            QScrollBar::add-line:vertical, QScrollBar::sub-line:vertical { height: 0px; }
        """)
        main_layout.addWidget(self.list_widget)
        main_layout.addSpacing(5)

        btn_layout = QHBoxLayout()
        btn_layout.setSpacing(12)
        btn_layout.setContentsMargins(0, 0, 0, 0)
        
        self.btn_load = create_styled_button("Abrir Pasta", "#8B00FF")
        self.btn_load.clicked.connect(self.load_folder)
        btn_layout.addWidget(self.btn_load)

        self.btn_obs = create_styled_button("Adicionar/editar\nobservação", "#228B22")
        self.btn_obs.clicked.connect(self.add_obs)
        btn_layout.addWidget(self.btn_obs)

        self.btn_collection = create_styled_button("Agrupar Filmes", "#FF6600")
        self.btn_collection.clicked.connect(self.open_collection)
        btn_layout.addWidget(self.btn_collection)

        self.btn_word = create_styled_button("Exportar Word", "#1E90FF")
        self.btn_word.clicked.connect(self.generate_word)
        btn_layout.addWidget(self.btn_word)

        self.btn_pdf = create_styled_button("Exportar PDF", "#FF0000")
        self.btn_pdf.clicked.connect(self.generate_pdf)
        btn_layout.addWidget(self.btn_pdf)

        main_layout.addLayout(btn_layout)
        self.setLayout(main_layout)

        self.folder_path = None
        self.folder_name = None
        self.files = []
        self.obs_data = {}
        self.order = []

    def base_dir(self):
        return os.path.dirname(sys.executable) if getattr(sys, 'frozen', False) else os.path.dirname(os.path.abspath(__file__))

    def json_path(self):
        return os.path.join(self.base_dir(), f"{self.folder_name}.json") if self.folder_name else None

    def xml_path(self):
        return os.path.join(self.base_dir(), f"{self.folder_name}.xml") if self.folder_name else None

    def load_folder(self):
        folder = QFileDialog.getExistingDirectory(self)
        if not folder: return
        self.folder_path = folder
        self.folder_name = os.path.basename(folder)
        self.files = sorted([f for f in os.listdir(folder) if os.path.isfile(os.path.join(folder, f)) and not is_hidden(os.path.join(folder, f))])
        self.load_obs()
        self.load_order()
        self.refresh(self.order)

    def load_obs(self):
        path = self.json_path()
        if path and os.path.exists(path):
            try:
                with open(path, encoding="utf-8") as f: self.obs_data = json.load(f).get("observacoes", {})
            except (json.JSONDecodeError, IOError): self.obs_data = {}
        else: self.obs_data = {}
        for f in self.files:
            if f not in self.obs_data: self.obs_data[f] = DEFAULT_OBS
        self.obs_data = {k: v for k, v in self.obs_data.items() if k in self.files}

    def save_obs(self):
        path = self.json_path()
        if path:
            with open(path, "w", encoding="utf-8") as f:
                json.dump({"pasta": self.folder_name, "observacoes": self.obs_data}, f, indent=4, ensure_ascii=False)

    def load_order(self):
        path = self.xml_path()
        if path and os.path.exists(path):
            try:
                tree = ET.parse(path)
                self.order = [e.text for e in tree.getroot().findall("file")]
            except ET.ParseError: self.order = []
        else: self.order = []
        for f in self.files:
            if f not in self.order: self.order.append(f)
        self.order = [f for f in self.order if f in self.files]

    def save_order(self):
        path = self.xml_path()
        if path:
            root = ET.Element("lista", pasta=self.folder_name)
            for f in self.order: ET.SubElement(root, "file").text = f
            ET.ElementTree(root).write(path, encoding="utf-8", xml_declaration=True)

    def refresh(self, base):
        self.list_widget.clear()
        for f in base:
            idx = self.order.index(f) + 1
            self.list_widget.addItem(f"{str(idx).zfill(4)}  - {os.path.splitext(f)[0]}")

    def filter_main(self):
        term = self.search_main.text().lower()
        self.refresh(self.order) if not term else self.refresh([f for f in self.order if term in f.lower()])

    def handle_delete_key(self, event):
        if event.key() == Qt.Key.Key_Delete:
            idx = self.list_widget.currentRow()
            if idx < 0: return
            term = self.search_main.text().lower()
            filtered = [f for f in self.order if term in f.lower()] if term else self.order
            if idx >= len(filtered): return
            f = filtered[idx]
            if f in self.order: self.order.remove(f)
            if f in self.obs_data: del self.obs_data[f]
            self.save_order()
            self.save_obs()
            self.refresh(self.order)
        else:
            QListWidget.keyPressEvent(self.list_widget, event)

    def add_obs(self):
        idx = self.list_widget.currentRow()
        if idx < 0: return
        term = self.search_main.text().lower()
        filtered = [f for f in self.order if term in f.lower()] if term else self.order
        if idx >= len(filtered): return
        f = filtered[idx]
        text, ok = QInputDialog.getMultiLineText(self, "OBS", f, self.obs_data[f])
        if ok:
            self.obs_data[f] = text if text.strip() else DEFAULT_OBS
            self.save_obs()

    def open_collection(self):
        dlg = CollectionDialog(self, self.order)
        if dlg.exec():
            new_seq = dlg.get_new_order()
            if not new_seq: return
            base_pos = min([self.order.index(f) for f in new_seq])
            for f in new_seq: self.order.remove(f)
            for i, f in enumerate(new_seq): self.order.insert(base_pos + i, f)
            self.save_order()
            self.refresh(self.order)

    def get_size(self, path):
        size = os.path.getsize(path)
        gb = size / (1024**3)
        return f"{round(gb,2)} GB" if gb >= 1 else f"{round(size/(1024**2),2)} MB"

    def set_bg(self, cell, color):
        tcPr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:fill'), color)
        tcPr.append(shd)

    def generate_word(self):
        path, _ = QFileDialog.getSaveFileName(self, "Salvar Word", "lista.docx", "Word Files (*.docx)")
        if not path: return
        doc = Document()
        section = doc.sections[0]
        section.top_margin = section.bottom_margin = section.left_margin = section.right_margin = Pt(10)
        table = doc.add_table(rows=0, cols=1)
        for i, f in enumerate(self.order, 1):
            row = table.add_row()
            row._tr.get_or_add_trPr().append(OxmlElement('w:cantSplit'))
            cell = row.cells[0]
            p = cell.paragraphs[0]
            name = os.path.splitext(f)[0]
            ext = os.path.splitext(f)[1]
            size = self.get_size(os.path.join(self.folder_path, f))
            run = p.add_run(f"{str(i).zfill(4)}  {name}\n")
            run.bold = True
            run.font.size = Pt(10)
            p.add_run("Informações Adicionais | ")
            p.add_run(f"Tamanho: {size} | ").font.color.rgb = RGBColor(255, 0, 0)
            p.add_run(f"Extensão {ext}\n").font.color.rgb = RGBColor(148, 0, 211)
            p.add_run(f"OBS: {self.obs_data[f]}").font.color.rgb = RGBColor(74, 163, 255)
            self.set_bg(cell, "EDEDED" if i % 2 == 0 else "FFFFFF")
        doc.save(path)

    def generate_pdf(self):
        path, _ = QFileDialog.getSaveFileName(self, "Salvar PDF", "lista.pdf", "PDF Files (*.pdf)")
        if not path: return

        c = canvas.Canvas(path, pagesize=A4)
        width, height = A4
        margin_left = 1.5 * cm
        margin_right = width - 1.5 * cm
        margin_top = 2 * cm
        margin_bottom = 2 * cm
        content_width = margin_right - margin_left

        c.setFont("Helvetica-Bold", 14)
        c.setFillColor(black)
        c.drawString(margin_left, height - margin_top, f"Lista de Arquivos - {self.folder_name}")
        c.setStrokeColor(gray)
        c.setLineWidth(1)
        c.line(margin_left, height - margin_top - 0.8 * cm, margin_right, height - margin_top - 0.8 * cm)
        
        y_position = height - margin_top - 1.5 * cm

        def estimate_block_height(title, obs):
            title_lines = self.wrap_text(title, content_width - 0.5 * cm, c, "Helvetica-Bold", 10)
            obs_lines = self.wrap_text(f"OBS: {obs}", content_width - 0.5 * cm, c, "Helvetica-Oblique", 8)
            return (len(title_lines) * 0.35) + 0.8 + (len(obs_lines) * 0.3) + 0.6  # em cm

        for i, f in enumerate(self.order, 1):
            name = os.path.splitext(f)[0]
            ext = os.path.splitext(f)[1]
            size = self.get_size(os.path.join(self.folder_path, f))
            obs = self.obs_data[f]

            block_height_cm = estimate_block_height(name, obs)
            if y_position - (block_height_cm * cm) < margin_bottom:
                c.showPage()
                y_position = height - margin_top - 1.5 * cm

            block_start_y = y_position
            title_lines = self.wrap_text(name, content_width - 0.5 * cm, c, "Helvetica-Bold", 10)
            c.setFont("Helvetica-Bold", 10)
            c.setFillColor(black)
            for line in title_lines:
                c.drawString(margin_left + 0.2 * cm, y_position, line)
                y_position -= 0.35 * cm

            c.setFont("Helvetica", 8)
            y_position -= 0.15 * cm
            c.setFillColor(HexColor("#FF0000"))
            c.drawString(margin_left + 0.2 * cm, y_position, f"Tamanho: {size}")
            c.setFillColor(HexColor("#9400D3"))
            c.drawString(margin_left + 5 * cm, y_position, f"Extensão: {ext}")

            y_position -= 0.35 * cm
            c.setFillColor(HexColor("#4AA3FF"))
            c.setFont("Helvetica-Oblique", 8)
            obs_lines = self.wrap_text(f"OBS: {obs}", content_width - 0.5 * cm, c, "Helvetica-Oblique", 8)
            for line in obs_lines:
                c.drawString(margin_left + 0.2 * cm, y_position, line)
                y_position -= 0.3 * cm
            y_position -= 0.2 * cm
            y_position -= 0.15 * cm

        c.save()

    def wrap_text(self, text, max_width, canvas_obj, font_name="Helvetica", font_size=10):
        words = text.split()
        lines = []
        current_line = ""
        canvas_obj.setFont(font_name, font_size)
        for word in words:
            test_line = f"{current_line} {word}".strip() if current_line else word
            if canvas_obj.stringWidth(test_line, font_name, font_size) <= max_width:
                current_line = test_line
            else:
                if current_line: lines.append(current_line)
                current_line = word
        if current_line: lines.append(current_line)
        return lines if lines else [""]


if __name__ == "__main__":
    import ctypes

    ctypes.windll.shell32.SetCurrentProcessExplicitAppUserModelID("anderson92tvrips")

    app = QApplication([])

    print(resource_path("icone.ico"))

    app.setWindowIcon(QIcon(resource_path("icone.ico")))

    show_license_popup()

    w = App()
    w.show()
    app.exec()
